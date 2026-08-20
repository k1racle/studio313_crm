from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


RUSSIAN_MONTHS = [
    'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря',
]


def format_money(value):
    number = f'{value:,.2f}'.replace(',', ' ').replace('.', ',')
    return number[:-3] if number.endswith(',00') else number


def format_document_date(value):
    return f'«{value.day:02d}» {RUSSIAN_MONTHS[value.month - 1]} {value.year} г.'


def set_paragraph_text(paragraph, text, *, bold=False, size=14):
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold


def get_template_path():
    configured = getattr(settings, 'PAYMENT_MEMO_TEMPLATE', '')
    candidates = [
        Path(configured) if configured else None,
        Path(settings.BASE_DIR).parent / 'Служебная записка.docx',
        Path(settings.BASE_DIR) / 'Служебная записка.docx',
    ]
    return next((path for path in candidates if path and path.exists()), None)


def create_fallback_document():
    document = Document()
    recipient = document.add_paragraph('Адресат')
    recipient.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    author_position = document.add_paragraph('От: сотрудник')
    author_position.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    author_name = document.add_paragraph('Ответственное лицо')
    author_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title = document.add_paragraph('Служебная записка')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    document.add_paragraph('Прошу согласовать денежные средства.')
    document.add_paragraph('Детализация суммы в рублях')
    document.add_paragraph('ИТОГО: 0 ₽.')
    date = document.add_paragraph('«01» января 2000 г.')
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature = document.add_paragraph('_________________________/Ответственное лицо/')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return document


def build_payment_memo(occurrence):
    plan = occurrence.plan
    author = plan.created_by
    template_path = get_template_path()
    document = Document(str(template_path)) if template_path else create_fallback_document()

    author_name = author.get_short_name() if author else 'Ответственное лицо'
    author_position = author.position if author and author.position else 'сотрудник'
    amount = format_money(occurrence.amount)
    generated_date = timezone.localdate()

    if document.tables:
        header_paragraphs = document.tables[0].cell(0, 1).paragraphs
        if len(header_paragraphs) >= 1:
            set_paragraph_text(header_paragraphs[0], plan.memo_recipient)
        if len(header_paragraphs) >= 2:
            set_paragraph_text(header_paragraphs[1], f'От: {author_position}')
        if len(header_paragraphs) >= 3:
            set_paragraph_text(header_paragraphs[2], author_name)
    else:
        set_paragraph_text(document.paragraphs[0], plan.memo_recipient)
        document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(document.paragraphs[1], f'От: {author_position}')
        document.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_text(document.paragraphs[2], author_name)
        document.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '№____' in paragraph.text:
                        set_paragraph_text(paragraph, f'№ ПЛ-{occurrence.id:05d}')
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    request_text = f'Прошу согласовать денежные средства на {plan.purpose.strip().rstrip(".")}.'
    details_text = (
        f'{plan.counterparty}. Срок оплаты — {occurrence.due_date.strftime("%d.%m.%Y")}. '
        f'Периодичность: {plan.get_frequency_display().lower()}. Сумма — {amount} рублей.'
    )

    replaced = {'request': False, 'details': False, 'total': False, 'date': False, 'signature': False}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith('Прошу'):
            set_paragraph_text(paragraph, request_text)
            replaced['request'] = True
        elif text.startswith('ИТОГО'):
            set_paragraph_text(paragraph, f'ИТОГО: {amount} ₽.', bold=True)
            replaced['total'] = True
        elif text.startswith('«') and ('г.' in text or '20' in text):
            set_paragraph_text(paragraph, format_document_date(generated_date))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            replaced['date'] = True
        elif '_________________________/' in text:
            set_paragraph_text(paragraph, f'_________________________/{author_name}/')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            replaced['signature'] = True
        elif text and not replaced['details'] and ('рублей' in text or 'клауд' in text.lower()):
            set_paragraph_text(paragraph, details_text)
            replaced['details'] = True

    if not replaced['request']:
        set_paragraph_text(document.paragraphs[2], request_text)
    if not replaced['details']:
        set_paragraph_text(document.paragraphs[3], details_text)
    if not replaced['total']:
        set_paragraph_text(document.paragraphs[4], f'ИТОГО: {amount} ₽.', bold=True)
    if not replaced['date']:
        set_paragraph_text(document.paragraphs[-2], format_document_date(generated_date))
    if not replaced['signature']:
        set_paragraph_text(document.paragraphs[-1], f'_________________________/{author_name}/')

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
